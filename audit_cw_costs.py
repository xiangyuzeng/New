#!/usr/bin/env python3
"""
CloudWatch Full-Dimension Cost Audit Tool — Luckin Coffee North America (LCNA)

Investigates all CloudWatch billing dimensions: log ingestion, custom metrics,
EC2 detailed monitoring, vended log pricing, alarms, and dashboards.
Generates reports with optimization recommendations ranked by savings.

Read-only — no changes are made. Remediation commands are output to reports only.

Minimum IAM Permissions Required:
{
  "Version": "2012-10-17",
  "Statement": [
    {
      "Effect": "Allow",
      "Action": [
        "logs:DescribeLogGroups",
        "cloudwatch:ListMetrics",
        "cloudwatch:GetMetricStatistics",
        "cloudwatch:DescribeAlarms",
        "cloudwatch:ListDashboards",
        "ec2:DescribeInstances",
        "ce:GetCostAndUsage",
        "sts:GetCallerIdentity"
      ],
      "Resource": "*"
    }
  ]
}

Usage:
    python audit_cw_costs.py                              # Full audit, markdown report
    python audit_cw_costs.py --format both                # Markdown + JSON + DOCX
    python audit_cw_costs.py --format docx                # DOCX only
    python audit_cw_costs.py --days 7 --skip-cost-explorer
    python audit_cw_costs.py --verbose --output-dir /tmp
"""

import argparse
import json
import logging
import os
import sys
import time
from datetime import datetime, timedelta, timezone
from typing import Any

import boto3
from botocore.exceptions import ClientError, BotoCoreError

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
# Log ingestion pricing (us-east-1)
INGESTION_COST_PER_GB = 0.50       # Standard log class
INGESTION_COST_PER_GB_IA = 0.25    # Infrequent Access log class
STORAGE_COST_PER_GB = 0.03         # Per GB stored per month

# Custom metrics pricing
METRIC_COST_FIRST_10K = 0.30       # First 10,000 metrics per month
METRIC_COST_NEXT_240K = 0.10       # Next 240,000
METRIC_COST_NEXT_750K = 0.05       # Next 750,000

# EC2 detailed monitoring: 7 metrics x $0.30 = $2.10/instance/month
EC2_DETAILED_METRICS_COUNT = 7
EC2_DETAILED_COST_PER_INSTANCE = EC2_DETAILED_METRICS_COUNT * METRIC_COST_FIRST_10K

# Alarms pricing
ALARM_COST_STANDARD = 0.10         # Standard resolution
ALARM_COST_HIGH_RES = 0.30         # High resolution (≤60s period)

# Dashboards pricing
DASHBOARD_FREE_TIER = 3
DASHBOARD_COST_EACH = 3.00         # Per dashboard per month beyond free tier

# API rate limiting
API_SLEEP_SECONDS = 0.1            # Sleep between GetMetricStatistics calls

# Zombie metric threshold
ZOMBIE_DAYS = 14                   # No data in this many days = zombie

logger = logging.getLogger("cw-cost-audit")


# ---------------------------------------------------------------------------
# Categorisation (aligned with audit_cw_logs.py)
# ---------------------------------------------------------------------------

def categorize(name: str) -> str:
    """Classify a log group name into a category."""
    n = name.lower()

    if "/aws/rds/" in n:
        if "slowquery" in n:
            return "RDS SlowQuery"
        if "error" in n:
            return "RDS Error"
        if "general" in n:
            return "RDS General"
        if "audit" in n:
            return "RDS Audit"
        return "RDS General"

    if "/aws/lambda/" in n:
        return "Lambda"

    if "/aws/apigateway/" in n or "api-gateway-execution-logs" in n:
        return "API Gateway"

    if "internetmonitor" in n:
        return "Internet Monitor"

    if "/aws/vpc/flowlogs" in n or "vpcflowlogs" in n:
        return "VPC Flow Logs"

    if "/ecs/" in n or "/aws/ecs/" in n:
        return "ECS / Container"

    if "/aws/elasticache/" in n:
        return "ElastiCache"

    if "aws-cloudtrail-logs" in n or "/aws/cloudtrail/" in n or "cloudtrail" in n:
        return "CloudTrail"

    if "/ec2/" in n:
        return "EC2 / Application"

    return "Other"


# ---------------------------------------------------------------------------
# AWS helpers
# ---------------------------------------------------------------------------

def get_session(region: str, profile: str | None) -> boto3.Session:
    kwargs: dict[str, Any] = {"region_name": region}
    if profile:
        kwargs["profile_name"] = profile
    return boto3.Session(**kwargs)


def get_account_id(session: boto3.Session) -> str:
    try:
        return session.client("sts").get_caller_identity()["Account"]
    except (ClientError, BotoCoreError) as exc:
        logger.warning("Could not determine account ID: %s", exc)
        return "unknown"


# ---------------------------------------------------------------------------
# Formatters
# ---------------------------------------------------------------------------

def _fmt_usd(val: float) -> str:
    return f"${val:,.2f}"


def _fmt_gb(val: float) -> str:
    return f"{val:,.2f}"


def _fmt_pct(val: float) -> str:
    return f"{val:.1f}%"


# ---------------------------------------------------------------------------
# TASK 1: Log Ingestion Volume Analysis
# ---------------------------------------------------------------------------

def audit_log_ingestion(
    logs_client,
    cw_client,
    days: int,
) -> dict[str, Any]:
    """Query IncomingBytes for every log group over the analysis window."""
    logger.info("TASK 1: Auditing log ingestion volumes (%d-day window)...", days)

    end_time = datetime.now(timezone.utc)
    start_time = end_time - timedelta(days=days)

    # Step 1: Discover all log groups
    log_groups: list[dict] = []
    paginator = logs_client.get_paginator("describe_log_groups")
    try:
        for page in paginator.paginate():
            for lg in page.get("logGroups", []):
                log_groups.append(lg)
    except ClientError as exc:
        logger.error("Failed to list log groups: %s", exc)
        return {"error": str(exc), "groups": [], "totalScanned": 0}

    logger.info("  Found %d log groups, querying IncomingBytes...", len(log_groups))

    # Step 2: Query IncomingBytes for each log group
    results: list[dict] = []
    for i, lg in enumerate(log_groups):
        name = lg["logGroupName"]
        log_class = lg.get("logGroupClass", "STANDARD")

        if (i + 1) % 25 == 0 or (i + 1) == len(log_groups):
            logger.info("  Progress: %d/%d log groups", i + 1, len(log_groups))

        try:
            resp = cw_client.get_metric_statistics(
                Namespace="AWS/Logs",
                MetricName="IncomingBytes",
                Dimensions=[{"Name": "LogGroupName", "Value": name}],
                StartTime=start_time,
                EndTime=end_time,
                Period=days * 86400,  # Single period spanning entire window
                Statistics=["Sum"],
            )
        except ClientError as exc:
            logger.debug("  Skipping %s: %s", name, exc)
            time.sleep(API_SLEEP_SECONDS)
            continue

        datapoints = resp.get("Datapoints", [])
        total_bytes = sum(dp.get("Sum", 0) for dp in datapoints)
        ingested_gb = total_bytes / (1024 ** 3)

        # Scale to monthly estimate if window != 30 days
        monthly_gb = ingested_gb * (30.0 / max(days, 1))
        unit_cost = INGESTION_COST_PER_GB_IA if log_class == "INFREQUENT_ACCESS" else INGESTION_COST_PER_GB
        monthly_cost = monthly_gb * unit_cost

        cat = categorize(name)

        # Recommendation logic
        recommendation = "OK"
        estimated_savings = 0.0
        if log_class == "STANDARD" and monthly_gb > 0:
            if cat in ("VPC Flow Logs",):
                recommendation = "Consider S3 direct delivery"
                estimated_savings = monthly_cost  # Eliminate CW ingestion entirely
            elif cat in ("RDS General",):
                recommendation = "Evaluate disabling general log"
                estimated_savings = monthly_cost
            elif monthly_gb >= 0.5:
                recommendation = "Candidate for Infrequent Access log class"
                estimated_savings = monthly_gb * (INGESTION_COST_PER_GB - INGESTION_COST_PER_GB_IA)

        results.append({
            "logGroupName": name,
            "category": cat,
            "logClass": log_class,
            "ingestedGB": round(ingested_gb, 4),
            "monthlyIngestedGB": round(monthly_gb, 4),
            "monthlyCost": round(monthly_cost, 4),
            "recommendation": recommendation,
            "estimatedSavings": round(estimated_savings, 4),
        })

        time.sleep(API_SLEEP_SECONDS)

    # Sort by cost descending
    results.sort(key=lambda r: r["monthlyCost"], reverse=True)

    # Aggregate by category
    by_cat: dict[str, dict] = {}
    for r in results:
        cat = r["category"]
        if cat not in by_cat:
            by_cat[cat] = {"count": 0, "monthlyGB": 0.0, "monthlyCost": 0.0, "savings": 0.0}
        by_cat[cat]["count"] += 1
        by_cat[cat]["monthlyGB"] += r["monthlyIngestedGB"]
        by_cat[cat]["monthlyCost"] += r["monthlyCost"]
        by_cat[cat]["savings"] += r["estimatedSavings"]

    total_monthly_cost = sum(r["monthlyCost"] for r in results)
    total_savings = sum(r["estimatedSavings"] for r in results)

    return {
        "totalScanned": len(log_groups),
        "totalWithIngestion": sum(1 for r in results if r["monthlyIngestedGB"] > 0),
        "totalMonthlyGB": round(sum(r["monthlyIngestedGB"] for r in results), 2),
        "totalMonthlyCost": round(total_monthly_cost, 2),
        "totalEstimatedSavings": round(total_savings, 2),
        "byCategory": {k: {kk: round(vv, 2) if isinstance(vv, float) else vv
                           for kk, vv in v.items()} for k, v in by_cat.items()},
        "groups": results,
    }


# ---------------------------------------------------------------------------
# TASK 2: Custom Metrics Audit
# ---------------------------------------------------------------------------

def audit_custom_metrics(
    cw_client,
    days: int,
) -> dict[str, Any]:
    """List all custom (non-AWS/) metrics, check for zombies and high cardinality."""
    logger.info("TASK 2: Auditing custom metrics...")

    end_time = datetime.now(timezone.utc)
    start_time = end_time - timedelta(days=ZOMBIE_DAYS)

    # Paginate list_metrics to get all metrics
    all_metrics: list[dict] = []
    try:
        paginator = cw_client.get_paginator("list_metrics")
        for page in paginator.paginate():
            for m in page.get("Metrics", []):
                all_metrics.append(m)
    except ClientError as exc:
        logger.error("Failed to list metrics: %s", exc)
        return {"error": str(exc), "namespaces": {}, "totalMetrics": 0}

    logger.info("  Total metrics discovered: %d", len(all_metrics))

    # Separate AWS-managed vs custom
    custom_metrics: list[dict] = []
    aws_metrics_count = 0
    for m in all_metrics:
        ns = m.get("Namespace", "")
        if ns.startswith("AWS/"):
            aws_metrics_count += 1
        else:
            custom_metrics.append(m)

    logger.info("  AWS-managed metrics: %d, Custom metrics: %d", aws_metrics_count, len(custom_metrics))

    # Group by namespace
    by_ns: dict[str, dict] = {}
    for m in custom_metrics:
        ns = m["Namespace"]
        if ns not in by_ns:
            by_ns[ns] = {
                "metrics": [],
                "metricNames": set(),
                "dimensions": {},  # dim_name -> set of values
            }
        by_ns[ns]["metrics"].append(m)
        by_ns[ns]["metricNames"].add(m["MetricName"])
        for dim in m.get("Dimensions", []):
            dn = dim["Name"]
            dv = dim["Value"]
            by_ns[ns]["dimensions"].setdefault(dn, set()).add(dv)

    # Build namespace summary
    ns_summaries: list[dict] = []
    total_custom = len(custom_metrics)

    # Calculate tiered cost
    def _tiered_cost(count: int) -> float:
        cost = 0.0
        remaining = count
        if remaining > 0:
            tier1 = min(remaining, 10000)
            cost += tier1 * METRIC_COST_FIRST_10K
            remaining -= tier1
        if remaining > 0:
            tier2 = min(remaining, 240000)
            cost += tier2 * METRIC_COST_NEXT_240K
            remaining -= tier2
        if remaining > 0:
            cost += remaining * METRIC_COST_NEXT_750K
        return cost

    total_cost = _tiered_cost(total_custom)

    # Check a sample of metrics in each namespace for zombie status
    zombie_check_count = 0
    for ns, data in sorted(by_ns.items(), key=lambda x: len(x[1]["metrics"]), reverse=True):
        metric_count = len(data["metrics"])
        unique_names = len(data["metricNames"])

        # Dimension cardinality analysis
        high_card_dims: list[str] = []
        for dn, values in data["dimensions"].items():
            if len(values) > 20:
                high_card_dims.append(f"{dn}({len(values)} values)")

        # Spot-check up to 5 metrics for zombie status
        active = 0
        zombie = 0
        checked = 0
        sample = data["metrics"][:5]
        for m in sample:
            if zombie_check_count >= 50:
                break  # Cap total zombie checks
            try:
                resp = cw_client.get_metric_statistics(
                    Namespace=ns,
                    MetricName=m["MetricName"],
                    Dimensions=m.get("Dimensions", []),
                    StartTime=start_time,
                    EndTime=end_time,
                    Period=ZOMBIE_DAYS * 86400,
                    Statistics=["Sum"],
                )
                if resp.get("Datapoints"):
                    active += 1
                else:
                    zombie += 1
                checked += 1
                zombie_check_count += 1
            except ClientError:
                pass
            time.sleep(API_SLEEP_SECONDS)

        # Extrapolate zombie rate
        zombie_rate = (zombie / checked * 100) if checked > 0 else 0
        est_zombie = int(metric_count * zombie / checked) if checked > 0 else 0

        ns_summaries.append({
            "namespace": ns,
            "metricCount": metric_count,
            "uniqueMetricNames": unique_names,
            "highCardinalityDimensions": high_card_dims,
            "sampledActive": active,
            "sampledZombie": zombie,
            "sampledTotal": checked,
            "estimatedZombieRate": round(zombie_rate, 1),
            "estimatedZombieCount": est_zombie,
            "recommendation": "Review high-cardinality dimensions" if high_card_dims else "OK",
        })

    ns_summaries.sort(key=lambda x: x["metricCount"], reverse=True)

    return {
        "totalAWSMetrics": aws_metrics_count,
        "totalCustomMetrics": total_custom,
        "estimatedMonthlyCost": round(total_cost, 2),
        "namespaceCount": len(by_ns),
        "namespaces": ns_summaries,
    }


# ---------------------------------------------------------------------------
# TASK 3: EC2 Detailed Monitoring Audit
# ---------------------------------------------------------------------------

def audit_ec2_monitoring(
    ec2_client,
) -> dict[str, Any]:
    """Check which EC2 instances have detailed monitoring enabled."""
    logger.info("TASK 3: Auditing EC2 detailed monitoring...")

    instances: list[dict] = []
    try:
        paginator = ec2_client.get_paginator("describe_instances")
        for page in paginator.paginate():
            for res in page.get("Reservations", []):
                for inst in res.get("Instances", []):
                    state = inst.get("State", {}).get("Name", "")
                    if state == "terminated":
                        continue

                    inst_id = inst["InstanceId"]
                    monitoring = inst.get("Monitoring", {}).get("State", "disabled")

                    # Get Name tag
                    name_tag = ""
                    for tag in inst.get("Tags", []):
                        if tag["Key"] == "Name":
                            name_tag = tag["Value"]
                            break

                    inst_type = inst.get("InstanceType", "")

                    instances.append({
                        "instanceId": inst_id,
                        "name": name_tag,
                        "instanceType": inst_type,
                        "state": state,
                        "monitoringState": monitoring,
                        "isDetailed": monitoring == "enabled",
                        "monthlyCost": round(EC2_DETAILED_COST_PER_INSTANCE, 2) if monitoring == "enabled" else 0.0,
                    })
    except ClientError as exc:
        logger.error("Failed to describe EC2 instances: %s", exc)
        return {"error": str(exc), "instances": [], "totalInstances": 0}

    detailed = [i for i in instances if i["isDetailed"]]
    basic = [i for i in instances if not i["isDetailed"]]
    total_cost = sum(i["monthlyCost"] for i in instances)

    # Generate disable commands for detailed instances
    disable_commands: list[str] = []
    for i in detailed:
        disable_commands.append(
            f"aws ec2 unmonitor-instances --instance-ids {i['instanceId']}  "
            f"# {i['name'] or 'unnamed'}"
        )

    logger.info("  Total: %d instances, Detailed: %d, Basic: %d",
                len(instances), len(detailed), len(basic))

    return {
        "totalInstances": len(instances),
        "detailedCount": len(detailed),
        "basicCount": len(basic),
        "detailedMonthlyCost": round(total_cost, 2),
        "estimatedSavings": round(total_cost, 2),  # All savings if switched to basic
        "disableCommands": disable_commands,
        "instances": instances,
    }


# ---------------------------------------------------------------------------
# TASK 4: Vended Log Pricing Verification
# ---------------------------------------------------------------------------

def audit_vended_logs(
    logs_client,
    ce_client,
    skip_ce: bool,
) -> dict[str, Any]:
    """Check Lambda/VPC Flow Log groups for log class and vended pricing."""
    logger.info("TASK 4: Verifying vended log pricing...")

    # Discover Lambda and VPC Flow Log groups
    vended_groups: list[dict] = []
    paginator = logs_client.get_paginator("describe_log_groups")
    try:
        for page in paginator.paginate():
            for lg in page.get("logGroups", []):
                name = lg["logGroupName"]
                cat = categorize(name)
                if cat in ("Lambda", "VPC Flow Logs"):
                    vended_groups.append({
                        "logGroupName": name,
                        "category": cat,
                        "logClass": lg.get("logGroupClass", "STANDARD"),
                        "storedBytes": lg.get("storedBytes", 0),
                    })
    except ClientError as exc:
        logger.error("Failed to list log groups for vended audit: %s", exc)

    logger.info("  Found %d vended-eligible log groups", len(vended_groups))

    # Cost Explorer breakdown
    ce_breakdown: list[dict] = []
    if not skip_ce and ce_client:
        try:
            end_date = datetime.now(timezone.utc).strftime("%Y-%m-01")
            # Go back 3 months
            start_dt = datetime.now(timezone.utc).replace(day=1) - timedelta(days=90)
            start_date = start_dt.strftime("%Y-%m-01")

            resp = ce_client.get_cost_and_usage(
                TimePeriod={"Start": start_date, "End": end_date},
                Granularity="MONTHLY",
                Metrics=["UnblendedCost"],
                Filter={
                    "Dimensions": {
                        "Key": "SERVICE",
                        "Values": ["AmazonCloudWatch"],
                    }
                },
                GroupBy=[{"Type": "DIMENSION", "Key": "USAGE_TYPE"}],
            )
            for period in resp.get("ResultsByTime", []):
                month = period["TimePeriod"]["Start"]
                for group in period.get("Groups", []):
                    usage_type = group["Keys"][0]
                    amount = float(group["Metrics"]["UnblendedCost"]["Amount"])
                    if amount > 0.001:
                        ce_breakdown.append({
                            "month": month,
                            "usageType": usage_type,
                            "cost": round(amount, 4),
                        })
            logger.info("  Cost Explorer returned %d usage-type rows", len(ce_breakdown))
        except ClientError as exc:
            logger.warning("  Cost Explorer query failed (may lack permissions): %s", exc)
            ce_breakdown = [{"error": str(exc)}]

    return {
        "vendedGroups": vended_groups,
        "lambdaCount": sum(1 for g in vended_groups if g["category"] == "Lambda"),
        "vpcFlowCount": sum(1 for g in vended_groups if g["category"] == "VPC Flow Logs"),
        "iaCount": sum(1 for g in vended_groups if g["logClass"] == "INFREQUENT_ACCESS"),
        "standardCount": sum(1 for g in vended_groups if g["logClass"] == "STANDARD"),
        "costExplorerBreakdown": ce_breakdown,
    }


# ---------------------------------------------------------------------------
# TASK 5: CloudWatch Alarms Audit
# ---------------------------------------------------------------------------

def audit_alarms(
    cw_client,
) -> dict[str, Any]:
    """Audit all CloudWatch alarms for waste and redundancy."""
    logger.info("TASK 5: Auditing CloudWatch alarms...")

    alarms: list[dict] = []
    try:
        paginator = cw_client.get_paginator("describe_alarms")
        for page in paginator.paginate():
            for a in page.get("MetricAlarms", []):
                period = a.get("Period", 300)
                is_high_res = period < 60
                cost = ALARM_COST_HIGH_RES if is_high_res else ALARM_COST_STANDARD

                alarms.append({
                    "alarmName": a["AlarmName"],
                    "namespace": a.get("Namespace", ""),
                    "metricName": a.get("MetricName", ""),
                    "state": a.get("StateValue", ""),
                    "period": period,
                    "isHighRes": is_high_res,
                    "monthlyCost": cost,
                    "dimensions": [
                        f"{d['Name']}={d['Value']}"
                        for d in a.get("Dimensions", [])
                    ],
                    "threshold": a.get("Threshold"),
                    "comparisonOperator": a.get("ComparisonOperator", ""),
                    "statistic": a.get("Statistic", ""),
                })
            for a in page.get("CompositeAlarms", []):
                alarms.append({
                    "alarmName": a["AlarmName"],
                    "namespace": "COMPOSITE",
                    "metricName": "",
                    "state": a.get("StateValue", ""),
                    "period": 0,
                    "isHighRes": False,
                    "monthlyCost": ALARM_COST_STANDARD,
                    "dimensions": [],
                    "threshold": None,
                    "comparisonOperator": "",
                    "statistic": "",
                })
    except ClientError as exc:
        logger.error("Failed to describe alarms: %s", exc)
        return {"error": str(exc), "alarms": [], "totalAlarms": 0}

    # Analyze
    by_state: dict[str, int] = {}
    for a in alarms:
        by_state[a["state"]] = by_state.get(a["state"], 0) + 1

    # Find INSUFFICIENT_DATA (likely dead targets)
    insufficient = [a for a in alarms if a["state"] == "INSUFFICIENT_DATA"]

    # Find duplicates (same namespace + metricName + dimensions + threshold)
    seen: dict[str, list[str]] = {}
    duplicates: list[list[str]] = []
    for a in alarms:
        if a["namespace"] == "COMPOSITE":
            continue
        key = f"{a['namespace']}|{a['metricName']}|{'|'.join(sorted(a['dimensions']))}|{a['threshold']}|{a['comparisonOperator']}"
        seen.setdefault(key, []).append(a["alarmName"])
    for key, names in seen.items():
        if len(names) > 1:
            duplicates.append(names)

    total_cost = sum(a["monthlyCost"] for a in alarms)
    insufficient_cost = sum(a["monthlyCost"] for a in insufficient)
    duplicate_cost = sum(
        sum(1 for _ in names[1:]) * ALARM_COST_STANDARD
        for names in duplicates
    )

    logger.info("  Total alarms: %d, INSUFFICIENT_DATA: %d, Duplicate sets: %d",
                len(alarms), len(insufficient), len(duplicates))

    return {
        "totalAlarms": len(alarms),
        "byState": by_state,
        "totalMonthlyCost": round(total_cost, 2),
        "insufficientDataCount": len(insufficient),
        "insufficientDataCost": round(insufficient_cost, 2),
        "insufficientDataAlarms": [a["alarmName"] for a in insufficient],
        "duplicateSets": duplicates,
        "duplicateSavings": round(duplicate_cost, 2),
        "estimatedSavings": round(insufficient_cost + duplicate_cost, 2),
        "alarms": alarms,
    }


# ---------------------------------------------------------------------------
# TASK 6: Dashboard Audit
# ---------------------------------------------------------------------------

def audit_dashboards(
    cw_client,
) -> dict[str, Any]:
    """List dashboards and calculate cost beyond free tier."""
    logger.info("TASK 6: Auditing CloudWatch dashboards...")

    dashboards: list[dict] = []
    try:
        paginator = cw_client.get_paginator("list_dashboards")
        for page in paginator.paginate():
            for d in page.get("DashboardEntries", []):
                dashboards.append({
                    "dashboardName": d["DashboardName"],
                    "dashboardArn": d.get("DashboardArn", ""),
                    "lastModified": d.get("LastModified", "").isoformat()
                    if hasattr(d.get("LastModified", ""), "isoformat")
                    else str(d.get("LastModified", "")),
                    "size": d.get("Size", 0),
                })
    except ClientError as exc:
        logger.error("Failed to list dashboards: %s", exc)
        return {"error": str(exc), "dashboards": [], "totalDashboards": 0}

    total = len(dashboards)
    billable = max(total - DASHBOARD_FREE_TIER, 0)
    cost = billable * DASHBOARD_COST_EACH

    logger.info("  Total dashboards: %d, Billable (beyond %d free): %d",
                total, DASHBOARD_FREE_TIER, billable)

    return {
        "totalDashboards": total,
        "freeTier": DASHBOARD_FREE_TIER,
        "billableCount": billable,
        "monthlyCost": round(cost, 2),
        "dashboards": dashboards,
    }


# ---------------------------------------------------------------------------
# Report generation — Markdown
# ---------------------------------------------------------------------------

def generate_markdown_report(
    account_id: str,
    region: str,
    days: int,
    ingestion: dict,
    custom_metrics: dict,
    ec2_monitoring: dict,
    vended_logs: dict,
    alarms_data: dict,
    dashboards_data: dict,
) -> str:
    now_str = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    # Grand totals
    total_monthly = (
        ingestion.get("totalMonthlyCost", 0)
        + custom_metrics.get("estimatedMonthlyCost", 0)
        + ec2_monitoring.get("detailedMonthlyCost", 0)
        + alarms_data.get("totalMonthlyCost", 0)
        + dashboards_data.get("monthlyCost", 0)
    )
    total_savings = (
        ingestion.get("totalEstimatedSavings", 0)
        + ec2_monitoring.get("estimatedSavings", 0)
        + alarms_data.get("estimatedSavings", 0)
    )

    lines: list[str] = []
    a = lines.append

    a("# CloudWatch Full-Dimension Cost Audit Report")
    a(f"**Date:** {now_str} | **Account:** {account_id} | **Region:** {region}")
    a(f"**Analysis Window:** {days} days")
    a("")

    # -- Executive Summary --
    a("## Executive Summary")
    a("")
    a(f"- **Estimated CloudWatch Monthly Cost:** {_fmt_usd(total_monthly)}")
    a(f"- **Estimated Annual Cost:** {_fmt_usd(total_monthly * 12)}")
    a(f"- **Total Achievable Monthly Savings:** {_fmt_usd(total_savings)}")
    a(f"- **Total Achievable Annual Savings:** {_fmt_usd(total_savings * 12)}")
    a("")
    a("### Cost Breakdown by Dimension")
    a("")
    a("| Dimension | Monthly Cost | % of Total |")
    a("|-----------|-------------|------------|")
    dimensions = [
        ("Log Ingestion", ingestion.get("totalMonthlyCost", 0)),
        ("Custom Metrics", custom_metrics.get("estimatedMonthlyCost", 0)),
        ("EC2 Detailed Monitoring", ec2_monitoring.get("detailedMonthlyCost", 0)),
        ("Alarms", alarms_data.get("totalMonthlyCost", 0)),
        ("Dashboards", dashboards_data.get("monthlyCost", 0)),
    ]
    for dim_name, dim_cost in dimensions:
        pct = (dim_cost / total_monthly * 100) if total_monthly > 0 else 0
        a(f"| {dim_name} | {_fmt_usd(dim_cost)} | {_fmt_pct(pct)} |")
    a(f"| **Total** | **{_fmt_usd(total_monthly)}** | **100.0%** |")
    a("")

    # -- Section 1: Log Ingestion --
    a("## 1. Log Ingestion Cost Analysis")
    a("")
    a(f"- **Log Groups Scanned:** {ingestion.get('totalScanned', 0)}")
    a(f"- **Groups with Ingestion:** {ingestion.get('totalWithIngestion', 0)}")
    a(f"- **Total Monthly Ingestion:** {_fmt_gb(ingestion.get('totalMonthlyGB', 0))} GB")
    a(f"- **Monthly Ingestion Cost:** {_fmt_usd(ingestion.get('totalMonthlyCost', 0))}")
    a(f"- **Estimated Savings:** {_fmt_usd(ingestion.get('totalEstimatedSavings', 0))}")
    a("")

    # Top 30 by ingestion
    a("### Top 30 Log Groups by Ingestion Volume")
    a("")
    a("| # | Log Group | Category | Ingested GB/mo | Cost/mo | Log Class | Recommendation | Savings/mo |")
    a("|---|-----------|----------|----------------|---------|-----------|----------------|------------|")
    top30 = ingestion.get("groups", [])[:30]
    for i, g in enumerate(top30, 1):
        name = g["logGroupName"]
        if len(name) > 55:
            name = name[:52] + "..."
        a(f"| {i} | {name} | {g['category']} | {_fmt_gb(g['monthlyIngestedGB'])} "
          f"| {_fmt_usd(g['monthlyCost'])} | {g['logClass']} "
          f"| {g['recommendation']} | {_fmt_usd(g['estimatedSavings'])} |")
    a("")

    # Category summary
    a("### Summary by Category")
    a("")
    a("| Category | Groups | Monthly GB | Monthly Cost |")
    a("|----------|--------|-----------|--------------|")
    by_cat = ingestion.get("byCategory", {})
    for cat in sorted(by_cat.keys()):
        c = by_cat[cat]
        a(f"| {cat} | {c['count']} | {_fmt_gb(c['monthlyGB'])} | {_fmt_usd(c['monthlyCost'])} |")
    a("")

    # -- Section 2: Custom Metrics --
    a("## 2. Custom Metrics Analysis")
    a("")
    a(f"- **AWS-Managed Metrics:** {custom_metrics.get('totalAWSMetrics', 0)}")
    a(f"- **Custom Metrics:** {custom_metrics.get('totalCustomMetrics', 0)}")
    a(f"- **Custom Namespaces:** {custom_metrics.get('namespaceCount', 0)}")
    a(f"- **Estimated Monthly Cost:** {_fmt_usd(custom_metrics.get('estimatedMonthlyCost', 0))}")
    a("")

    a("### Summary by Namespace")
    a("")
    a("| Namespace | Metrics | Unique Names | Zombie Rate | High-Cardinality Dims | Recommendation |")
    a("|-----------|---------|--------------|-------------|----------------------|----------------|")
    for ns in custom_metrics.get("namespaces", []):
        hc = ", ".join(ns["highCardinalityDimensions"]) if ns["highCardinalityDimensions"] else "—"
        a(f"| {ns['namespace']} | {ns['metricCount']} | {ns['uniqueMetricNames']} "
          f"| {_fmt_pct(ns['estimatedZombieRate'])} | {hc} | {ns['recommendation']} |")
    a("")

    # -- Section 3: EC2 Detailed Monitoring --
    a("## 3. EC2 Detailed Monitoring")
    a("")
    a(f"- **Total Instances:** {ec2_monitoring.get('totalInstances', 0)}")
    a(f"- **Detailed Monitoring Enabled:** {ec2_monitoring.get('detailedCount', 0)}")
    a(f"- **Basic Monitoring:** {ec2_monitoring.get('basicCount', 0)}")
    a(f"- **Monthly Detailed Monitoring Cost:** {_fmt_usd(ec2_monitoring.get('detailedMonthlyCost', 0))}")
    a(f"- **Potential Savings (switch all to basic):** {_fmt_usd(ec2_monitoring.get('estimatedSavings', 0))}")
    a("")

    if ec2_monitoring.get("detailedCount", 0) > 0:
        a("### Instances with Detailed Monitoring")
        a("")
        a("| Instance ID | Name | Type | State | Cost/mo |")
        a("|------------|------|------|-------|---------|")
        for inst in ec2_monitoring.get("instances", []):
            if inst["isDetailed"]:
                name = inst["name"] or "—"
                if len(name) > 35:
                    name = name[:32] + "..."
                a(f"| {inst['instanceId']} | {name} | {inst['instanceType']} "
                  f"| {inst['state']} | {_fmt_usd(inst['monthlyCost'])} |")
        a("")

        a("### Disable Commands")
        a("")
        a("```bash")
        for cmd in ec2_monitoring.get("disableCommands", [])[:20]:
            a(cmd)
        if len(ec2_monitoring.get("disableCommands", [])) > 20:
            a(f"# ... and {len(ec2_monitoring['disableCommands']) - 20} more")
        a("```")
        a("")

    # -- Section 4: Vended Logs --
    a("## 4. Vended Log Pricing Verification")
    a("")
    a(f"- **Lambda Log Groups:** {vended_logs.get('lambdaCount', 0)}")
    a(f"- **VPC Flow Log Groups:** {vended_logs.get('vpcFlowCount', 0)}")
    a(f"- **Infrequent Access Class:** {vended_logs.get('iaCount', 0)}")
    a(f"- **Standard Class:** {vended_logs.get('standardCount', 0)}")
    a("")

    ce_data = vended_logs.get("costExplorerBreakdown", [])
    if ce_data and not (len(ce_data) == 1 and "error" in ce_data[0]):
        a("### CloudWatch Cost Explorer Breakdown (Last 3 Months)")
        a("")
        a("| Month | Usage Type | Cost |")
        a("|-------|-----------|------|")
        for row in ce_data:
            if "error" not in row:
                a(f"| {row['month']} | {row['usageType']} | {_fmt_usd(row['cost'])} |")
        a("")
    elif ce_data and "error" in ce_data[0]:
        a(f"> **Note:** Cost Explorer query skipped or failed: {ce_data[0].get('error', 'unknown')}")
        a("")

    # -- Section 5: Alarms --
    a("## 5. Alarms Audit")
    a("")
    a(f"- **Total Alarms:** {alarms_data.get('totalAlarms', 0)}")
    a(f"- **Monthly Alarm Cost:** {_fmt_usd(alarms_data.get('totalMonthlyCost', 0))}")
    a(f"- **INSUFFICIENT_DATA Alarms:** {alarms_data.get('insufficientDataCount', 0)} "
      f"({_fmt_usd(alarms_data.get('insufficientDataCost', 0))}/mo)")
    a(f"- **Duplicate Alarm Sets:** {len(alarms_data.get('duplicateSets', []))}")
    a(f"- **Estimated Savings:** {_fmt_usd(alarms_data.get('estimatedSavings', 0))}")
    a("")

    by_state = alarms_data.get("byState", {})
    if by_state:
        a("### Alarms by State")
        a("")
        a("| State | Count |")
        a("|-------|-------|")
        for state, count in sorted(by_state.items()):
            a(f"| {state} | {count} |")
        a("")

    insuff = alarms_data.get("insufficientDataAlarms", [])
    if insuff:
        a("### INSUFFICIENT_DATA Alarms (likely dead targets)")
        a("")
        for alarm_name in insuff[:30]:
            a(f"- `{alarm_name}`")
        if len(insuff) > 30:
            a(f"- *... and {len(insuff) - 30} more*")
        a("")

    dups = alarms_data.get("duplicateSets", [])
    if dups:
        a("### Duplicate Alarm Sets")
        a("")
        for dup_set in dups[:10]:
            a(f"- {', '.join(f'`{n}`' for n in dup_set)}")
        if len(dups) > 10:
            a(f"- *... and {len(dups) - 10} more sets*")
        a("")

    # -- Section 6: Dashboards --
    a("## 6. Dashboard Audit")
    a("")
    a(f"- **Total Dashboards:** {dashboards_data.get('totalDashboards', 0)}")
    a(f"- **Free Tier:** {DASHBOARD_FREE_TIER}")
    a(f"- **Billable Dashboards:** {dashboards_data.get('billableCount', 0)}")
    a(f"- **Monthly Cost:** {_fmt_usd(dashboards_data.get('monthlyCost', 0))}")
    a("")

    if dashboards_data.get("dashboards"):
        a("| Dashboard | Last Modified | Size (bytes) |")
        a("|-----------|--------------|--------------|")
        for d in dashboards_data["dashboards"]:
            a(f"| {d['dashboardName']} | {d['lastModified']} | {d['size']} |")
        a("")

    # -- Section 7: Optimization Recommendations --
    a("## 7. Optimization Recommendations (Ranked by Savings)")
    a("")
    a("| # | Item | Current/mo | Savings/mo | Annual Savings | Risk | Action |")
    a("|---|------|-----------|------------|----------------|------|--------|")

    recs: list[tuple] = []
    if ec2_monitoring.get("estimatedSavings", 0) > 0:
        recs.append((
            "Disable EC2 detailed monitoring (Prometheus covers this)",
            ec2_monitoring["detailedMonthlyCost"],
            ec2_monitoring["estimatedSavings"],
            "Low — Prometheus already collecting at higher resolution",
            "`aws ec2 unmonitor-instances --instance-ids <ids>`",
        ))
    if ingestion.get("totalEstimatedSavings", 0) > 0:
        recs.append((
            "Switch high-volume log groups to Infrequent Access / S3",
            ingestion["totalMonthlyCost"],
            ingestion["totalEstimatedSavings"],
            "Medium — verify query patterns before switching to IA",
            "Update logGroupClass or redirect to S3",
        ))
    if alarms_data.get("estimatedSavings", 0) > 0:
        recs.append((
            "Delete INSUFFICIENT_DATA alarms + deduplicate",
            alarms_data["totalMonthlyCost"],
            alarms_data["estimatedSavings"],
            "Low — these alarms are non-functional",
            "`aws cloudwatch delete-alarms --alarm-names <names>`",
        ))
    if dashboards_data.get("monthlyCost", 0) > 0:
        recs.append((
            "Remove unused dashboards beyond free tier",
            dashboards_data["monthlyCost"],
            dashboards_data["monthlyCost"],
            "Low — verify dashboards are unused first",
            "`aws cloudwatch delete-dashboards --dashboard-names <names>`",
        ))

    recs.sort(key=lambda r: r[2], reverse=True)
    for i, (item, current, savings, risk, action) in enumerate(recs, 1):
        a(f"| {i} | {item} | {_fmt_usd(current)} | {_fmt_usd(savings)} "
          f"| {_fmt_usd(savings * 12)} | {risk} | {action} |")
    a("")

    # -- Footer --
    a("---")
    a(f"*Generated by `audit_cw_costs.py` on {now_str}*")

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Report generation — JSON
# ---------------------------------------------------------------------------

def generate_json_report(
    account_id: str,
    region: str,
    days: int,
    ingestion: dict,
    custom_metrics: dict,
    ec2_monitoring: dict,
    vended_logs: dict,
    alarms_data: dict,
    dashboards_data: dict,
) -> str:
    total_monthly = (
        ingestion.get("totalMonthlyCost", 0)
        + custom_metrics.get("estimatedMonthlyCost", 0)
        + ec2_monitoring.get("detailedMonthlyCost", 0)
        + alarms_data.get("totalMonthlyCost", 0)
        + dashboards_data.get("monthlyCost", 0)
    )
    total_savings = (
        ingestion.get("totalEstimatedSavings", 0)
        + ec2_monitoring.get("estimatedSavings", 0)
        + alarms_data.get("estimatedSavings", 0)
    )

    report = {
        "reportDate": datetime.now(timezone.utc).isoformat(),
        "accountId": account_id,
        "region": region,
        "analysisWindowDays": days,
        "estimatedMonthlyCost": round(total_monthly, 2),
        "estimatedAnnualCost": round(total_monthly * 12, 2),
        "estimatedMonthlySavings": round(total_savings, 2),
        "estimatedAnnualSavings": round(total_savings * 12, 2),
        "logIngestion": ingestion,
        "customMetrics": custom_metrics,
        "ec2DetailedMonitoring": ec2_monitoring,
        "vendedLogs": vended_logs,
        "alarms": alarms_data,
        "dashboards": dashboards_data,
    }
    return json.dumps(report, indent=2, default=str)


# ---------------------------------------------------------------------------
# Report generation — DOCX
# ---------------------------------------------------------------------------

def generate_docx_report(
    account_id: str,
    region: str,
    days: int,
    ingestion: dict,
    custom_metrics: dict,
    ec2_monitoring: dict,
    vended_logs: dict,
    alarms_data: dict,
    dashboards_data: dict,
    output_path: str,
) -> None:
    """Generate a professional Word document report."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    now_str = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    doc = Document()

    # -- Styles --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    def _add_table(headers: list[str], rows: list[list[str]]) -> None:
        """Add a formatted table to the document."""
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        # Header row
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9)
        # Data rows
        for row_data in rows:
            row = table.add_row()
            for i, val in enumerate(row_data):
                row.cells[i].text = str(val)
                for p in row.cells[i].paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
        doc.add_paragraph()  # spacing

    # Grand totals
    total_monthly = (
        ingestion.get("totalMonthlyCost", 0)
        + custom_metrics.get("estimatedMonthlyCost", 0)
        + ec2_monitoring.get("detailedMonthlyCost", 0)
        + alarms_data.get("totalMonthlyCost", 0)
        + dashboards_data.get("monthlyCost", 0)
    )
    total_savings = (
        ingestion.get("totalEstimatedSavings", 0)
        + ec2_monitoring.get("estimatedSavings", 0)
        + alarms_data.get("estimatedSavings", 0)
    )

    # -- Title --
    title = doc.add_heading("CloudWatch Full-Dimension Cost Audit Report", level=0)
    p = doc.add_paragraph()
    p.add_run(f"Date: ").bold = True
    p.add_run(f"{now_str}  |  ")
    p.add_run(f"Account: ").bold = True
    p.add_run(f"{account_id}  |  ")
    p.add_run(f"Region: ").bold = True
    p.add_run(f"{region}  |  ")
    p.add_run(f"Window: ").bold = True
    p.add_run(f"{days} days")

    # -- Executive Summary --
    doc.add_heading("Executive Summary", level=1)
    p = doc.add_paragraph()
    for label, val in [
        ("Estimated Monthly Cost", _fmt_usd(total_monthly)),
        ("Estimated Annual Cost", _fmt_usd(total_monthly * 12)),
        ("Achievable Monthly Savings", _fmt_usd(total_savings)),
        ("Achievable Annual Savings", _fmt_usd(total_savings * 12)),
    ]:
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(f"{val}\n")

    _add_table(
        ["Dimension", "Monthly Cost", "% of Total"],
        [
            [name, _fmt_usd(cost), _fmt_pct(cost / total_monthly * 100 if total_monthly else 0)]
            for name, cost in [
                ("Log Ingestion", ingestion.get("totalMonthlyCost", 0)),
                ("Custom Metrics", custom_metrics.get("estimatedMonthlyCost", 0)),
                ("EC2 Detailed Monitoring", ec2_monitoring.get("detailedMonthlyCost", 0)),
                ("Alarms", alarms_data.get("totalMonthlyCost", 0)),
                ("Dashboards", dashboards_data.get("monthlyCost", 0)),
            ]
        ] + [["Total", _fmt_usd(total_monthly), "100.0%"]],
    )

    # -- Section 1: Log Ingestion --
    doc.add_heading("1. Log Ingestion Cost Analysis", level=1)
    p = doc.add_paragraph()
    for label, val in [
        ("Log Groups Scanned", str(ingestion.get("totalScanned", 0))),
        ("Groups with Ingestion", str(ingestion.get("totalWithIngestion", 0))),
        ("Total Monthly Ingestion", f"{_fmt_gb(ingestion.get('totalMonthlyGB', 0))} GB"),
        ("Monthly Ingestion Cost", _fmt_usd(ingestion.get("totalMonthlyCost", 0))),
        ("Estimated Savings", _fmt_usd(ingestion.get("totalEstimatedSavings", 0))),
    ]:
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(f"{val}\n")

    doc.add_heading("Top 30 Log Groups by Ingestion Volume", level=2)
    top30 = ingestion.get("groups", [])[:30]
    _add_table(
        ["#", "Log Group", "Category", "GB/mo", "Cost/mo", "Class", "Recommendation", "Savings/mo"],
        [
            [
                str(i),
                g["logGroupName"][:55] + ("..." if len(g["logGroupName"]) > 55 else ""),
                g["category"],
                _fmt_gb(g["monthlyIngestedGB"]),
                _fmt_usd(g["monthlyCost"]),
                g["logClass"],
                g["recommendation"],
                _fmt_usd(g["estimatedSavings"]),
            ]
            for i, g in enumerate(top30, 1)
        ],
    )

    doc.add_heading("Summary by Category", level=2)
    by_cat = ingestion.get("byCategory", {})
    _add_table(
        ["Category", "Groups", "Monthly GB", "Monthly Cost"],
        [
            [cat, str(c["count"]), _fmt_gb(c["monthlyGB"]), _fmt_usd(c["monthlyCost"])]
            for cat, c in sorted(by_cat.items())
        ],
    )

    # -- Section 2: Custom Metrics --
    doc.add_heading("2. Custom Metrics Analysis", level=1)
    p = doc.add_paragraph()
    for label, val in [
        ("AWS-Managed Metrics", str(custom_metrics.get("totalAWSMetrics", 0))),
        ("Custom Metrics", str(custom_metrics.get("totalCustomMetrics", 0))),
        ("Custom Namespaces", str(custom_metrics.get("namespaceCount", 0))),
        ("Estimated Monthly Cost", _fmt_usd(custom_metrics.get("estimatedMonthlyCost", 0))),
    ]:
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(f"{val}\n")

    doc.add_heading("Summary by Namespace", level=2)
    _add_table(
        ["Namespace", "Metrics", "Unique Names", "Zombie Rate", "High-Card Dims", "Recommendation"],
        [
            [
                ns["namespace"],
                str(ns["metricCount"]),
                str(ns["uniqueMetricNames"]),
                _fmt_pct(ns["estimatedZombieRate"]),
                ", ".join(ns["highCardinalityDimensions"]) or "\u2014",
                ns["recommendation"],
            ]
            for ns in custom_metrics.get("namespaces", [])
        ],
    )

    # -- Section 3: EC2 Detailed Monitoring --
    doc.add_heading("3. EC2 Detailed Monitoring", level=1)
    p = doc.add_paragraph()
    for label, val in [
        ("Total Instances", str(ec2_monitoring.get("totalInstances", 0))),
        ("Detailed Monitoring Enabled", str(ec2_monitoring.get("detailedCount", 0))),
        ("Basic Monitoring", str(ec2_monitoring.get("basicCount", 0))),
        ("Monthly Cost", _fmt_usd(ec2_monitoring.get("detailedMonthlyCost", 0))),
        ("Potential Savings", _fmt_usd(ec2_monitoring.get("estimatedSavings", 0))),
    ]:
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(f"{val}\n")

    if ec2_monitoring.get("detailedCount", 0) > 0:
        doc.add_heading("Instances with Detailed Monitoring", level=2)
        _add_table(
            ["Instance ID", "Name", "Type", "State", "Cost/mo"],
            [
                [
                    inst["instanceId"],
                    (inst["name"] or "\u2014")[:35],
                    inst["instanceType"],
                    inst["state"],
                    _fmt_usd(inst["monthlyCost"]),
                ]
                for inst in ec2_monitoring.get("instances", [])
                if inst["isDetailed"]
            ],
        )

    # -- Section 4: Vended Logs --
    doc.add_heading("4. Vended Log Pricing Verification", level=1)
    p = doc.add_paragraph()
    for label, val in [
        ("Lambda Log Groups", str(vended_logs.get("lambdaCount", 0))),
        ("VPC Flow Log Groups", str(vended_logs.get("vpcFlowCount", 0))),
        ("Infrequent Access Class", str(vended_logs.get("iaCount", 0))),
        ("Standard Class", str(vended_logs.get("standardCount", 0))),
    ]:
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(f"{val}\n")

    ce_data = vended_logs.get("costExplorerBreakdown", [])
    if ce_data and not (len(ce_data) == 1 and "error" in ce_data[0]):
        doc.add_heading("CloudWatch Cost Explorer Breakdown (Last 3 Months)", level=2)
        _add_table(
            ["Month", "Usage Type", "Cost"],
            [[row["month"], row["usageType"], _fmt_usd(row["cost"])]
             for row in ce_data if "error" not in row],
        )

    # -- Section 5: Alarms --
    doc.add_heading("5. Alarms Audit", level=1)
    p = doc.add_paragraph()
    for label, val in [
        ("Total Alarms", str(alarms_data.get("totalAlarms", 0))),
        ("Monthly Alarm Cost", _fmt_usd(alarms_data.get("totalMonthlyCost", 0))),
        ("INSUFFICIENT_DATA Alarms", f"{alarms_data.get('insufficientDataCount', 0)} "
         f"({_fmt_usd(alarms_data.get('insufficientDataCost', 0))}/mo)"),
        ("Duplicate Alarm Sets", str(len(alarms_data.get("duplicateSets", [])))),
        ("Estimated Savings", _fmt_usd(alarms_data.get("estimatedSavings", 0))),
    ]:
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(f"{val}\n")

    by_state = alarms_data.get("byState", {})
    if by_state:
        _add_table(
            ["State", "Count"],
            [[state, str(count)] for state, count in sorted(by_state.items())],
        )

    insuff = alarms_data.get("insufficientDataAlarms", [])
    if insuff:
        doc.add_heading("INSUFFICIENT_DATA Alarms (likely dead targets)", level=2)
        for name in insuff[:30]:
            doc.add_paragraph(name, style="List Bullet")
        if len(insuff) > 30:
            doc.add_paragraph(f"... and {len(insuff) - 30} more", style="List Bullet")

    # -- Section 6: Dashboards --
    doc.add_heading("6. Dashboard Audit", level=1)
    p = doc.add_paragraph()
    for label, val in [
        ("Total Dashboards", str(dashboards_data.get("totalDashboards", 0))),
        ("Free Tier", str(DASHBOARD_FREE_TIER)),
        ("Billable Dashboards", str(dashboards_data.get("billableCount", 0))),
        ("Monthly Cost", _fmt_usd(dashboards_data.get("monthlyCost", 0))),
    ]:
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(f"{val}\n")

    if dashboards_data.get("dashboards"):
        _add_table(
            ["Dashboard", "Last Modified", "Size (bytes)"],
            [[d["dashboardName"], d["lastModified"], str(d["size"])]
             for d in dashboards_data["dashboards"]],
        )

    # -- Section 7: Optimization Recommendations --
    doc.add_heading("7. Optimization Recommendations (Ranked by Savings)", level=1)

    recs: list[tuple] = []
    if ec2_monitoring.get("estimatedSavings", 0) > 0:
        recs.append((
            "Disable EC2 detailed monitoring (Prometheus covers this)",
            ec2_monitoring["detailedMonthlyCost"],
            ec2_monitoring["estimatedSavings"],
            "Low",
            "aws ec2 unmonitor-instances --instance-ids <ids>",
        ))
    if ingestion.get("totalEstimatedSavings", 0) > 0:
        recs.append((
            "Switch high-volume log groups to Infrequent Access / S3",
            ingestion["totalMonthlyCost"],
            ingestion["totalEstimatedSavings"],
            "Medium",
            "Update logGroupClass or redirect to S3",
        ))
    if alarms_data.get("estimatedSavings", 0) > 0:
        recs.append((
            "Delete INSUFFICIENT_DATA alarms + deduplicate",
            alarms_data["totalMonthlyCost"],
            alarms_data["estimatedSavings"],
            "Low",
            "aws cloudwatch delete-alarms --alarm-names <names>",
        ))
    if dashboards_data.get("monthlyCost", 0) > 0:
        recs.append((
            "Remove unused dashboards beyond free tier",
            dashboards_data["monthlyCost"],
            dashboards_data["monthlyCost"],
            "Low",
            "aws cloudwatch delete-dashboards --dashboard-names <names>",
        ))

    recs.sort(key=lambda r: r[2], reverse=True)
    _add_table(
        ["#", "Item", "Current/mo", "Savings/mo", "Annual Savings", "Risk", "Action"],
        [
            [str(i), item, _fmt_usd(current), _fmt_usd(savings),
             _fmt_usd(savings * 12), risk, action]
            for i, (item, current, savings, risk, action) in enumerate(recs, 1)
        ],
    )

    # -- Footer --
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f"Generated by audit_cw_costs.py on {now_str}")
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.save(output_path)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="CloudWatch Full-Dimension Cost Audit Tool",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s                                    # Full audit, markdown report
  %(prog)s --format both                      # Markdown + JSON
  %(prog)s --days 7 --skip-cost-explorer      # 7-day window, skip CE
  %(prog)s --verbose --output-dir /tmp         # Verbose to /tmp
        """,
    )
    parser.add_argument("--region", default="us-east-1", help="AWS region (default: us-east-1)")
    parser.add_argument("--profile", default=None, help="AWS CLI profile name")
    parser.add_argument("--format", choices=["markdown", "json", "docx", "both"], default="markdown",
                        help="Output format (default: markdown). 'both' generates all three.")
    parser.add_argument("--output-dir", default="./reports", help="Output directory (default: ./reports)")
    parser.add_argument("--days", type=int, default=30, help="Analysis window in days (default: 30)")
    parser.add_argument("--skip-cost-explorer", action="store_true",
                        help="Skip Cost Explorer queries (saves $0.01/call)")
    parser.add_argument("--verbose", action="store_true", help="Enable verbose logging")
    return parser.parse_args(argv)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv)

    # Logging
    log_level = logging.DEBUG if args.verbose else logging.INFO
    logging.basicConfig(
        level=log_level,
        format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
    )

    logger.info("Starting CloudWatch Full-Dimension Cost Audit")
    logger.info("Region: %s | Days: %d | Skip CE: %s",
                args.region, args.days, args.skip_cost_explorer)

    # AWS session
    try:
        session = get_session(args.region, args.profile)
        account_id = get_account_id(session)
        logs_client = session.client("logs")
        cw_client = session.client("cloudwatch")
        ec2_client = session.client("ec2")
        ce_client = session.client("ce") if not args.skip_cost_explorer else None
    except Exception as exc:
        logger.critical("Failed to initialize AWS session: %s", exc)
        return 2

    logger.info("Account: %s", account_id)

    # Run all 6 investigation tasks
    try:
        ingestion = audit_log_ingestion(logs_client, cw_client, args.days)
    except Exception as exc:
        logger.error("TASK 1 failed: %s", exc)
        ingestion = {"error": str(exc), "groups": [], "totalScanned": 0,
                     "totalMonthlyCost": 0, "totalEstimatedSavings": 0}

    try:
        custom_metrics = audit_custom_metrics(cw_client, args.days)
    except Exception as exc:
        logger.error("TASK 2 failed: %s", exc)
        custom_metrics = {"error": str(exc), "namespaces": [],
                          "totalCustomMetrics": 0, "estimatedMonthlyCost": 0}

    try:
        ec2_monitoring = audit_ec2_monitoring(ec2_client)
    except Exception as exc:
        logger.error("TASK 3 failed: %s", exc)
        ec2_monitoring = {"error": str(exc), "instances": [],
                          "totalInstances": 0, "detailedMonthlyCost": 0, "estimatedSavings": 0}

    try:
        vended_logs = audit_vended_logs(logs_client, ce_client, args.skip_cost_explorer)
    except Exception as exc:
        logger.error("TASK 4 failed: %s", exc)
        vended_logs = {"error": str(exc), "vendedGroups": []}

    try:
        alarms_data = audit_alarms(cw_client)
    except Exception as exc:
        logger.error("TASK 5 failed: %s", exc)
        alarms_data = {"error": str(exc), "alarms": [], "totalAlarms": 0,
                       "totalMonthlyCost": 0, "estimatedSavings": 0}

    try:
        dashboards_data = audit_dashboards(cw_client)
    except Exception as exc:
        logger.error("TASK 6 failed: %s", exc)
        dashboards_data = {"error": str(exc), "dashboards": [],
                           "totalDashboards": 0, "monthlyCost": 0}

    # Generate reports
    os.makedirs(args.output_dir, exist_ok=True)
    date_str = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    base_name = f"cloudwatch-cost-audit-{date_str}"
    outputs: list[str] = []

    report_args = (
        account_id, args.region, args.days,
        ingestion, custom_metrics, ec2_monitoring,
        vended_logs, alarms_data, dashboards_data,
    )

    if args.format in ("markdown", "both"):
        md = generate_markdown_report(*report_args)
        md_path = os.path.join(args.output_dir, f"{base_name}.md")
        with open(md_path, "w") as f:
            f.write(md)
        outputs.append(md_path)
        logger.info("Markdown report: %s", md_path)

    if args.format in ("json", "both"):
        js = generate_json_report(*report_args)
        json_path = os.path.join(args.output_dir, f"{base_name}.json")
        with open(json_path, "w") as f:
            f.write(js)
        outputs.append(json_path)
        logger.info("JSON report: %s", json_path)

    if args.format in ("docx", "both"):
        docx_path = os.path.join(args.output_dir, f"{base_name}.docx")
        generate_docx_report(*report_args, output_path=docx_path)
        outputs.append(docx_path)
        logger.info("DOCX report: %s", docx_path)

    # Summary to stdout
    total_monthly = (
        ingestion.get("totalMonthlyCost", 0)
        + custom_metrics.get("estimatedMonthlyCost", 0)
        + ec2_monitoring.get("detailedMonthlyCost", 0)
        + alarms_data.get("totalMonthlyCost", 0)
        + dashboards_data.get("monthlyCost", 0)
    )
    total_savings = (
        ingestion.get("totalEstimatedSavings", 0)
        + ec2_monitoring.get("estimatedSavings", 0)
        + alarms_data.get("estimatedSavings", 0)
    )

    print(f"\n{'='*60}")
    print(f"CloudWatch Full-Dimension Cost Audit — {date_str}")
    print(f"{'='*60}")
    print(f"  Account:              {account_id}")
    print(f"  Region:               {args.region}")
    print(f"  Analysis window:      {args.days} days")
    print(f"  Log groups scanned:   {ingestion.get('totalScanned', 0)}")
    print(f"  Custom metrics:       {custom_metrics.get('totalCustomMetrics', 0)}")
    print(f"  EC2 instances:        {ec2_monitoring.get('totalInstances', 0)}")
    print(f"  Alarms:               {alarms_data.get('totalAlarms', 0)}")
    print(f"  Dashboards:           {dashboards_data.get('totalDashboards', 0)}")
    print(f"  ─────────────────────────────────────────")
    print(f"  Est. monthly cost:    ${total_monthly:,.2f}")
    print(f"  Est. annual cost:     ${total_monthly * 12:,.2f}")
    print(f"  Est. savings/mo:      ${total_savings:,.2f}")
    print(f"  Est. savings/yr:      ${total_savings * 12:,.2f}")
    for p in outputs:
        print(f"  Report:               {p}")
    print(f"{'='*60}\n")

    return 0


if __name__ == "__main__":
    sys.exit(main())
